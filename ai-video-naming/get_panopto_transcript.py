#!/usr/bin/env python3
"""
Panopto Transcript Downloader
------------------------------
Downloads the transcript (VTT captions) for a given Panopto video session
and saves it as a .docx file named {recording_date}__{recording_name}.docx.

Usage:
    python get_panopto_transcript.py <session_url_or_id>
    python get_panopto_transcript.py <session_url_or_id> -o output.docx
"""

import requests
import json
import os
import re
import sys
import argparse
from datetime import datetime, timedelta
import base64
from urllib.parse import unquote
from dotenv import load_dotenv
from docx import Document

load_dotenv(override=True)

PANOPTO_SERVER = os.getenv("PANOPTO_SERVER", "ioe.cloud.panopto.eu")
PANOPTO_CLIENT_ID = os.getenv("PANOPTO_CLIENT_ID", "")
PANOPTO_CLIENT_SECRET = os.getenv("PANOPTO_CLIENT_SECRET", "")
TOKEN_FILE = os.getenv("TOKEN_FILE", "panopto_tokens.json")


def extract_session_id(url_or_id):
    """Extract session ID from a Panopto viewer URL or return as-is if already an ID"""
    if re.match(r'^[a-f0-9-]{36}$', url_or_id, re.IGNORECASE):
        return url_or_id

    # Viewer URL: .../Panopto/Pages/Viewer.aspx?id=<guid>
    match = re.search(r'[?&]id=([a-f0-9-]{36})', url_or_id, re.IGNORECASE)
    if match:
        return match.group(1)

    # Embed URL or other formats with session id
    match = re.search(r'[?&]sessionId=([a-f0-9-]{36})', url_or_id, re.IGNORECASE)
    if match:
        return match.group(1)

    return None


def get_auth_token():
    """Get a valid access token, refreshing if needed"""
    if not os.path.exists(TOKEN_FILE):
        print("No token file found. Run the main sync script first to authenticate.")
        return None

    with open(TOKEN_FILE, 'r') as f:
        token_data = json.load(f)

    try:
        expires_at = datetime.fromisoformat(token_data['expires_at'])
        if datetime.now() < expires_at:
            return token_data['access_token']
    except:
        pass

    # Refresh token
    token_url = f"https://{PANOPTO_SERVER}/Panopto/oauth2/connect/token"
    data = {'grant_type': 'refresh_token', 'refresh_token': token_data['refresh_token']}
    credentials = f"{PANOPTO_CLIENT_ID}:{PANOPTO_CLIENT_SECRET}"
    encoded_credentials = base64.b64encode(credentials.encode()).decode()
    headers = {'Content-Type': 'application/x-www-form-urlencoded', 'Authorization': f'Basic {encoded_credentials}'}

    response = requests.post(token_url, data=data, headers=headers, timeout=30)
    if response.status_code == 200:
        new_token_data = response.json()
        access_token = new_token_data['access_token']

        save_data = {
            'access_token': access_token,
            'refresh_token': new_token_data.get('refresh_token', token_data['refresh_token']),
            'expires_at': (datetime.now() + timedelta(seconds=new_token_data.get('expires_in', 3600) - 300)).isoformat(),
            'created_at': datetime.now().isoformat()
        }
        with open(TOKEN_FILE, 'w') as f:
            json.dump(save_data, f, indent=2)

        return access_token

    print(f"Token refresh failed: {response.status_code}")
    return None


def get_legacy_auth_cookie(auth_token):
    """Get legacy auth cookie for caption downloads"""
    url = f"https://{PANOPTO_SERVER}/Panopto/api/v1/auth/legacyLogin"
    headers = {"Authorization": f"Bearer {auth_token}"}

    response = requests.get(url, headers=headers, timeout=30)
    if response.status_code == 200 and '.ASPXAUTH' in response.cookies:
        return response.cookies
    return None


def get_session_details(auth_token, session_id):
    """Get full session details"""
    url = f"https://{PANOPTO_SERVER}/Panopto/api/v1/sessions/{session_id}"
    headers = {"Authorization": f"Bearer {auth_token}"}

    response = requests.get(url, headers=headers, timeout=30)
    if response.status_code == 200:
        return response.json()
    return None


def download_transcript(auth_token, cookies, session_id):
    """Download the VTT transcript for a session"""
    details = get_session_details(auth_token, session_id)
    if not details:
        print(f"Could not get session details for {session_id}")
        return None, None, None

    session_name = details.get("Name", "Unknown")
    start_time = details.get("StartTime")
    caption_url = details.get("Urls", {}).get("CaptionDownloadUrl")
    if not caption_url:
        print(f"No captions available for session: {session_name}")
        return None, session_name, start_time

    response = requests.get(caption_url, cookies=cookies, timeout=60)
    if response.status_code == 200:
        return response.text, session_name, start_time

    print(f"Failed to download captions (HTTP {response.status_code})")
    return None, session_name, start_time


def parse_vtt_to_text(vtt_content):
    """Parse VTT content and extract timestamped text entries"""
    lines = vtt_content.strip().splitlines()
    entries = []
    current_timestamp = None
    current_text = []

    for line in lines:
        line = line.strip()
        # Skip empty lines, WEBVTT header, NOTE lines
        if not line or line.startswith("WEBVTT") or line.startswith("NOTE"):
            if current_timestamp and current_text:
                entries.append((current_timestamp, ' '.join(current_text)))
                current_timestamp = None
                current_text = []
            continue
        # Skip cue identifiers (numeric only)
        if re.match(r'^\d+$', line):
            continue
        # Timestamp line like 00:00:01.234 --> 00:00:05.678
        ts_match = re.match(r'(\d{2}:\d{2}:\d{2}[.,]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[.,]\d{3})', line)
        if ts_match:
            if current_timestamp and current_text:
                entries.append((current_timestamp, ' '.join(current_text)))
                current_text = []
            start = ts_match.group(1).replace(',', '.').rsplit('.', 1)[0]
            current_timestamp = start
            continue
        # Text line
        current_text.append(line)

    if current_timestamp and current_text:
        entries.append((current_timestamp, ' '.join(current_text)))

    return entries


def sanitize_filename(name):
    """Remove characters that are invalid in Windows filenames"""
    return re.sub(r'[<>:"/\\|?*]', '_', name)


def save_transcript_as_docx(vtt_content, session_name, start_time, output_path=None):
    """Convert VTT transcript to a .docx file"""
    date_str = None

    # Try API StartTime first
    if start_time:
        try:
            dt = datetime.fromisoformat(start_time.replace("Z", "+00:00"))
            date_str = dt.strftime("%Y-%m-%d")
        except (ValueError, AttributeError):
            pass

    # Fall back to parsing date from session name like [30/03/2026 09:00 AM]
    if not date_str and session_name:
        m = re.search(r'\[(\d{2})/(\d{2})/(\d{4})', session_name)
        if m:
            date_str = f"{m.group(3)}-{m.group(2)}-{m.group(1)}"

    if not date_str:
        date_str = "unknown-date"

    if not output_path:
        # Strip leading date bracket like [30/03/2026 09:00 AM] from name
        clean_name = re.sub(r'^\[.*?\]\s*', '', session_name)
        safe_name = sanitize_filename(clean_name)
        os.makedirs("transcripts", exist_ok=True)
        output_path = os.path.join("transcripts", f"{date_str}__{safe_name}.docx")

    text_lines = parse_vtt_to_text(vtt_content)

    doc = Document()
    doc.add_heading(session_name, level=1)
    doc.add_paragraph(f"Recording date: {date_str}")
    doc.add_paragraph("")  # spacer

    for timestamp, text in text_lines:
        doc.add_paragraph(f"[{timestamp}]  {text}")

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Download Panopto transcript for a video session")
    parser.add_argument("session", help="Panopto session ID (GUID) or viewer URL")
    parser.add_argument("-o", "--output", help="Output .docx file path (default: {date}__{name}.docx)")
    args = parser.parse_args()

    session_id = extract_session_id(args.session)
    if not session_id:
        print(f"Could not extract session ID from: {args.session}")
        sys.exit(1)

    auth_token = get_auth_token()
    if not auth_token:
        sys.exit(1)

    cookies = get_legacy_auth_cookie(auth_token)
    if not cookies:
        print("Could not get legacy auth cookie for caption download")
        sys.exit(1)

    transcript, session_name, start_time = download_transcript(auth_token, cookies, session_id)
    if not transcript:
        sys.exit(1)

    output_path = save_transcript_as_docx(transcript, session_name, start_time, args.output)
    print(f"Transcript saved to: {output_path}")
    print(f"Session: {session_name}")


if __name__ == "__main__":
    main()
